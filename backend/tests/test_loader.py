"""Tests for core.loader (load_document, load_txt, load_docx, load_md)."""

import sys
from pathlib import Path

BACKEND_DIR = Path(__file__).resolve().parent.parent
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

import pytest
from core.loader import load_document, load_txt


class TestLoadTxt:
    def test_normal(self, tmp_txt_file):
        text = load_txt(tmp_txt_file)
        assert "测试内容" in text
        assert "Hello" in text

    def test_empty(self, tmp_empty_txt):
        text = load_txt(tmp_empty_txt)
        assert text == ""


class TestLoadDocument:
    def test_txt(self, tmp_txt_file):
        text = load_document(tmp_txt_file)
        assert "测试内容" in text

    def test_unsupported_extension(self, tmp_unsupported_file):
        with pytest.raises(ValueError, match="不支持的文件类型"):
            load_document(tmp_unsupported_file)

    def test_doc_extension(self, tmp_path):
        path = tmp_path / "old.doc"
        path.write_text("fake word doc", encoding="utf-8")
        with pytest.raises(ValueError, match="不支持 .doc 格式"):
            load_document(path)

    def test_md(self, tmp_path):
        path = tmp_path / "readme.md"
        path.write_text("# Hello\n\nThis is **bold** and ![image](img.png)", encoding="utf-8")
        text = load_document(path)
        assert "Hello" in text
        assert "![image]" not in text  # image syntax removed

    def test_docx_empty(self, tmp_path):
        """Minimal valid .docx should not crash."""
        from docx import Document as DocxDocument
        path = tmp_path / "empty.docx"
        doc = DocxDocument()
        doc.save(str(path))
        text = load_document(path)
        assert isinstance(text, str)

    def test_docx_with_text(self, tmp_path):
        from docx import Document as DocxDocument
        path = tmp_path / "hello.docx"
        doc = DocxDocument()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("中文测试")
        doc.save(str(path))
        text = load_document(path)
        assert "Hello World" in text
        assert "中文测试" in text
