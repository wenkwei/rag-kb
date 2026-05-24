import re
from pathlib import Path
from typing import Union, Optional

import fitz
import numpy as np
import pdfplumber
from paddleocr import PaddleOCR
from docx import Document as DocxDocument


_ocr: Optional[PaddleOCR] = None


def _get_ocr() -> PaddleOCR:
    """Lazy-init singleton PaddleOCR instance (heavy model ~100MB loaded once)."""
    global _ocr
    if _ocr is None:
        _ocr = PaddleOCR(use_angle_cls=True, lang="ch", use_gpu=False, show_log=False)
    return _ocr


def _ocr_page(file_path: Union[str, Path], page_num: int, fitz_doc: fitz.Document) -> str:
    """Render a PDF page to an image and run OCR to extract text."""
    try:
        ocr = _get_ocr()
        pix = fitz_doc[page_num].get_pixmap(dpi=200)
        img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
            pix.height, pix.width, pix.n
        )
        result = ocr.ocr(img_array, cls=True)
        if result and result[0]:
            return "\n".join(line[1][0] for line in result[0])
    except Exception:
        pass
    return ""


def load_pdf(file_path: Union[str, Path]) -> str:
    """Extract text from a PDF.

    Uses pdfplumber for fast text extraction (handles digital/text PDFs).
    Falls back to PaddleOCR on pages where pdfplumber finds little text
    (handles scanned / image-based PDFs). Per-page, the method that
    returns more content wins.
    """
    text_parts: list[str] = []
    fitz_doc: Optional[fitz.Document] = None

    try:
        with pdfplumber.open(str(file_path)) as pdf:
            fitz_doc = fitz.open(str(file_path))

            for page_num, page in enumerate(pdf.pages):
                pdfplumber_text = page.extract_text() or ""

                if len(pdfplumber_text.strip()) < 50:
                    ocr_text = _ocr_page(file_path, page_num, fitz_doc)
                    page_text = (
                        pdfplumber_text
                        if len(pdfplumber_text.strip()) >= len(ocr_text.strip())
                        else ocr_text
                    )
                else:
                    page_text = pdfplumber_text

                if page_text.strip():
                    text_parts.append(page_text)
    finally:
        if fitz_doc:
            fitz_doc.close()

    return "\n".join(text_parts)


def load_txt(file_path: Union[str, Path]) -> str:
    """Read text from a TXT file."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def load_docx(file_path: Union[str, Path]) -> str:
    """Extract text from a DOCX file using python-docx.

    Includes paragraph text and table content (pipe-separated).
    """
    doc = DocxDocument(str(file_path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n".join(parts)


def load_md(file_path: Union[str, Path]) -> str:
    """Extract plain text from a Markdown file.

    Removes image syntax ``![alt](url)`` and link URL parts ``[text](url)``,
    keeping the link text.
    """
    raw = Path(file_path).read_text("utf-8")
    # Remove images: ![alt](url)
    text = re.sub(r"!\[.*?\]\(.*?\)", "", raw)
    # Remove link URLs, keep text: [text](url) -> text
    text = re.sub(r"\[([^\]]*)\]\(.*?\)", r"\1", text)
    return text.strip()


def load_document(file_path: Union[str, Path]) -> str:
    """Load text from a supported document type (PDF, TXT, DOCX, or MD).

    Raises ValueError for unsupported file types.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return load_pdf(path)
    elif suffix == ".txt":
        return load_txt(path)
    elif suffix == ".docx":
        return load_docx(path)
    elif suffix == ".md":
        return load_md(path)
    elif suffix == ".doc":
        raise ValueError("不支持 .doc 格式（旧版 Word），请转换为 .docx 后重试")
    else:
        raise ValueError(f"不支持的文件类型: {suffix}，仅支持 PDF、TXT、DOCX 和 MD")
